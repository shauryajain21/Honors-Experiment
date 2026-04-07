#!/usr/bin/env python3
"""Convert analysis report and discussion to formatted .docx files"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUT = Path('/Users/shaurya/Honors-Experiment/thesis_analysis')
NYU_PURPLE = RGBColor(0x57, 0x06, 0x8C)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x66, 0x66, 0x66)

def style_heading(run, size=14, color=NYU_PURPLE, bold=True):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NYU_PURPLE

def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    run.bold = bold
    run.italic = italic
    return p

def add_stat(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    run.font.color.rgb = DARK
    return p

def add_figure(doc, path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    r.italic = True

# ============================================================
# ANALYSIS REPORT DOCX
# ============================================================
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Honors Research')
run.font.size = Pt(28)
run.font.color.rgb = NYU_PURPLE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Data Analysis Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x89, 0x00, 0xE1)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ball-and-Jar Bayesian Probability Updating Experiment')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('N = 23 participants  |  2,277 experiment trials')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_page_break()

# --- Read the report ---
report_text = open(OUT / 'analysis_report.txt').read()
sections = report_text.split('\n')

current_section = None
i = 0
while i < len(sections):
    line = sections[i]

    # Skip separator lines
    if line.startswith('===') or line.startswith('---'):
        i += 1
        continue

    # Section headers (numbered)
    if line and line[0].isdigit() and '. ' in line[:5]:
        add_styled_heading(doc, line, level=1)
        i += 1
        continue

    # Sub-content
    if line.strip().startswith('Prediction:') or line.strip().startswith('->'):
        add_body(doc, line.strip(), italic=True)
        i += 1
        continue

    # Stats lines (indented or with specific patterns)
    if line.strip() and (line.startswith('  ') or 'p =' in line or 'r =' in line or
                          't =' in line or 'Mean' in line[:10] or 'Phase' in line[:10] or
                          line.strip().startswith('|') or 'RMSE' in line[:10]):
        add_stat(doc, line.strip())
        i += 1
        continue

    # Table-like headers
    if 'Code' in line and 'RMSE' in line:
        add_stat(doc, line.strip())
        i += 1
        # Read table rows
        while i < len(sections) and sections[i].strip():
            add_stat(doc, sections[i].strip())
            i += 1
        continue

    # Regular text
    if line.strip():
        add_body(doc, line.strip())

    i += 1

# Add figures
doc.add_page_break()
add_styled_heading(doc, 'Figures', level=1)

figures = [
    ('figure1_summary.png', 'Figure 1: Six-panel analysis summary — error distribution, calibration, individual differences, phase effects, prior retrieval, and confidence.'),
    ('figure2_prior_retrieval.png', 'Figure 2: Prior retrieval analysis — retrieval vs continuation distances, Bayesian model comparison, and transition point magnitudes.'),
    ('figure3_temporal.png', 'Figure 3: Temporal trends across trials — reaction time, confidence, and absolute error by phase.'),
    ('figure4_conservatism.png', 'Figure 4: Conservatism analysis — scatter plot with regression line and calibration by Bayesian decile.'),
    ('figure5_individuals.png', 'Figure 5: Example individual participant trajectories with Bayesian-Reset and Bayesian-Retrieve overlays.'),
    ('figure6_correlations.png', 'Figure 6: Participant-level metric correlation matrix.'),
]

for fname, caption in figures:
    fpath = OUT / fname
    if fpath.exists():
        add_figure(doc, fpath, caption, width=6.5 if 'summary' in fname or 'individual' in fname else 6.0)
        doc.add_paragraph()

doc.save(OUT / 'Analysis_Report.docx')
print(f"Saved: {OUT / 'Analysis_Report.docx'}")

# ============================================================
# DISCUSSION DOCX
# ============================================================
doc2 = Document()
style = doc2.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Discussion')
run.font.size = Pt(24)
run.font.color.rgb = NYU_PURPLE
run.bold = True
doc2.add_paragraph()

discussion_text = open(OUT / 'discussion.txt').read()

for line in discussion_text.split('\n'):
    if not line.strip():
        continue

    if line.startswith('===') or line.startswith('DISCUSSION'):
        continue

    if line.startswith('---'):
        continue

    # Section headers
    if line.startswith('HYPOTHESIS') or line.startswith('ADDITIONAL') or \
       line.startswith('LIMITATIONS') or line.startswith('CONCLUSION'):
        add_styled_heading(doc2, line.strip(), level=2)
        continue

    # Body text
    p = doc2.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(line.strip())
    run.font.size = Pt(12)
    run.font.color.rgb = DARK

doc2.save(OUT / 'Discussion.docx')
print(f"Saved: {OUT / 'Discussion.docx'}")
print("\nDone!")
